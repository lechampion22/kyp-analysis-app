import streamlit as st
from datetime import datetime
import docx
from io import BytesIO
from docx.shared import Pt

# Force the page to use a wide layout
st.set_page_config(layout="wide")

# ---------------------------
# Pre-determined text for Equities funds
# ---------------------------
equities_texts = {
    "DFA Global Equity Portfolio F (DFA607)": (
        "Low Cost:\n"
        "Total cost is 0.32%. This is competitive, especially given the science-based, factor-driven strategy DFA employs.\n\n"
        "Diversification:\n"
        "Globally diversified across developed and emerging markets, emphasizing value, small-cap, and profitability factors (13,000+ stock holdings).\n\n"
        "Evidence-Based Philosophy:\n"
        "Strong evidence-based philosophy grounded in academic research, focusing on factors like value, size, and profitability.\n\n"
        "Efficient Trading:\n"
        "Employs patient and flexible trading strategies to reduce costs and optimize implementation.\n\n"
        "Track Record:\n"
        "Long history of delivering risk-adjusted returns aligned with its evidence-based strategy.\n\n"
        "Recommendations:\n"
        "For clients prioritizing low costs and an evidence-based philosophy, DFA Global Equity Portfolio F (DFA607) is ideal due to its unique factor tilts and efficient trading. Its strategy aligns with clients seeking robust diversification, academic rigor, and long-term value creation."
    ),
    "iShares Core Equity ETF Portfolio (XEQT)": (
        "Low Cost:\n"
        "Extremely low MER of 0.20%, making it a highly cost-effective option.\n\n"
        "Diversification:\n"
        "Provides global equity exposure with a market-cap-weighted strategy. Includes developed and emerging markets but lacks a factor-based tilt (8,000+ stock holding).\n\n"
        "Evidence-Based Philosophy:\n"
        "Evidence-based in terms of market efficiency but does not incorporate factor tilts.\n\n"
        "Efficient Trading:\n"
        "ETFs benefit from market liquidity and efficient execution, though rebalancing costs are borne by the fund.\n\n"
        "Track Record:\n"
        "Relatively new but backed by BlackRock's extensive ETF management experience.\n\n"
        "Recommendations:\n"
        "iShares Core Equity ETF Portfolio (XEQT) is excellent for clients seeking ultra-low costs and simplicity. Its minimal fees and market-cap-weighted approach make it a practical choice for cost-conscious investors who prefer passive strategies. While XEQT is not directly accessible, it serves as a benchmark for assessing the efficiency and cost-effectiveness of global equity solutions."
    ),
    "Mackenzie Core Fund Portfolio": (
        "Low Cost:\n"
        "MERs range around 1.26%, significantly higher due to active management.\n\n"
        "Diversification:\n"
        "Offers diversification through a mix of actively managed funds. However, potential overlap and manager bias may reduce true diversification benefits (100+ stock holdings).\n\n"
        "Evidence-Based Philosophy:\n"
        "Active management often deviates from evidence-based principles, relying on manager judgment rather than systematic factors.\n\n"
        "Efficient Trading:\n"
        "Active management may result in higher trading costs due to frequent portfolio adjustments.\n\n"
        "Track Record:\n"
        "Performance varies significantly by fund, often underperforming benchmarks after fees.\n\n"
        "Recommendations:\n"
        "The Mackenzie Core Fund Portfolio may appeal to clients who prefer active management and are willing to pay higher fees for the potential of outperformance. However, these clients should be made aware of the risks and the historical challenges active management has faced in consistently beating benchmarks."
    ),
    "RBC Global Equity Index ETF Fund (RBF2146)": (
        "Low Cost:\n"
        "MER of approximately 0.49%, competitive for an index-following strategy but higher than XEQT due to additional fees embedded in the mutual fund structure.\n\n"
        "Diversification:\n"
        "Tracks a globally diversified index but follows the same strategy as XEQT, lacking factor tilts.\n\n"
        "Evidence-Based Philosophy:\n"
        "Tracks indices, adhering to market efficiency but without factor integration.\n\n"
        "Efficient Trading:\n"
        "Efficient trading tied to index replication but lacks DFA's strategic trading nuances.\n\n"
        "Track Record:\n"
        "Performance closely tracks its benchmark, ensuring market returns minus fees.\n\n"
        "Recommendations:\n"
        "RBC Global Equity Index ETF Fund (RBF2146) serves as a middle ground, offering simplicity and a cost structure slightly higher than ETFs but without the advanced factor-based strategies of DFA. It is suitable for clients who value ease of use and a reliable index-tracking approach without needing to commit to a fully evidence-based investment philosophy."
    ),
    "Fidelity All-in-One Equity ETF F (FID7567)": (
        "Low Cost:\n"
        "MER of 0.59%, offering a balance of cost efficiency and active allocation flexibility within an all-in-one ETF structure.\n\n"
        "Diversification:\n"
        "A globally diversified portfolio with a modest tactical asset allocation approach, aiming to enhance diversification through periodic rebalancing and active shifts (1,500+ stock holdings).\n\n"
        "Evidence-Based Philosophy:\n"
        "Combines evidence-based passive management with dynamic rebalancing to adapt to market changes while maintaining a diversified core.\n\n"
        "Efficient Trading:\n"
        "Maintains efficient trading practices with periodic rebalancing aimed at optimizing portfolio alignment with market conditions.\n\n"
        "Track Record:\n"
        "Relatively new with limited historical data, but backed by Fidelity’s established expertise in portfolio management.\n\n"
        "Recommendations:\n"
        "Fidelity All-in-One Equity ETF F (FID7567) is suitable for clients looking for a cost-effective solution with some active asset allocation and rebalancing elements. Its slightly higher MER is justified by the flexibility and tactical adjustments it offers, making it a good choice for those who want a blend of passive management and active oversight. Additionally, its ability to be held in the client’s name makes it particularly advantageous for accounts like RESPs or situations where nominee fees are a concern."
    )
}

# ---------------------------
# Pre-determined text for Fixed Income funds
# ---------------------------
fixed_income_texts = {
    "DFA Global Fixed Income Portfolio F": (
        "Low Cost:\n"
        "Total cost is 0.31%. This is competitive given the strategic factor-driven approach.\n\n"
        "Diversification:\n"
        "Broad global diversification, including government, corporate, and inflation-protected bonds. The fund emphasizes credit and term premiums.\n\n"
        "Evidence-Based Philosophy:\n"
        "Grounded in academic research, with disciplined exposure to credit and term premiums while managing interest rate risk.\n\n"
        "Efficient Trading:\n"
        "Uses patient and flexible trading strategies to reduce costs and optimize implementation.\n\n"
        "Track Record:\n"
        "Long history of delivering risk-adjusted returns through systematic, globally diversified strategies.\n\n"
        "Recommendations:\n"
        "Ideal for clients seeking global diversification and an evidence-based approach to fixed income, emphasizing term and credit premiums."
    ),
    "iShares Core CAD Universe Bond Index ETF": (
        "Low Cost:\n"
        "MER of 0.10%, making it one of the most cost-efficient options for fixed-income exposure in Canada. RBF Fund MER = 0.16%.\n\n"
        "Diversification:\n"
        "Offers exposure to a wide array of Canadian bonds, including government and investment-grade corporate bonds.\n\n"
        "Evidence-Based Philosophy:\n"
        "Passive indexing aligns with evidence-based principles, tracking a broad Canadian bond market index.\n\n"
        "Efficient Trading:\n"
        "Benefits from the efficiency of ETF structures and liquidity in the Canadian bond market.\n\n"
        "Track Record:\n"
        "Well-established ETF with a consistent track record of tracking its benchmark effectively.\n\n"
        "Recommendations:\n"
        "Suitable for clients looking for ultra-low costs and broad exposure to the Canadian bond market."
    ),
    "Mackenzie Unconstrained Fund F": (
        "Low Cost:\n"
        "MER of 0.78%, significantly higher due to its actively managed, flexible strategy.\n\n"
        "Diversification:\n"
        "Highly flexible and diversified, with the ability to invest in global fixed-income opportunities, including high-yield and emerging market bonds.\n\n"
        "Evidence-Based Philosophy:\n"
        "Active management deviates from strict evidence-based approaches, relying on manager expertise and judgment.\n\n"
        "Efficient Trading:\n"
        "Active trading may result in higher transaction costs, though this is offset by the potential for higher returns in less liquid markets.\n\n"
        "Track Record:\n"
        "Performance varies based on market conditions and manager decisions, offering potential for higher returns but with greater risk.\n\n"
        "Recommendations:\n"
        "Best for clients who are comfortable with active management and the potential for higher risk and returns through flexible global strategies."
    ),
    "Fidelity Systematic Canadian Bond Index ETF": (
        "Low Cost:\n"
        "MER of 0.37%, providing a balance of cost efficiency and systematic indexing.\n\n"
        "Diversification:\n"
        "Focuses on the Canadian bond market, tracking a systematic strategy across government and corporate bonds.\n\n"
        "Evidence-Based Philosophy:\n"
        "Combines evidence-based principles with systematic indexing, adhering to a disciplined investment process.\n\n"
        "Efficient Trading:\n"
        "Efficiently tracks its benchmark while maintaining low costs.\n\n"
        "Track Record:\n"
        "Relatively new but backed by Fidelity’s expertise in systematic fixed-income strategies.\n\n"
        "Recommendations:\n"
        "A balanced choice for clients seeking low costs with a systematic, evidence-based approach focused on the Canadian market."
    ),
    "Lysander-Canso Corporate Value Bond F": (
        "Low Cost:\n"
        "MER of 0.90% (high), reflecting active management with a focus on corporate bonds.\n\n"
        "Diversification:\n"
        "Concentrated on Canadian corporate bonds, offering a unique diversification element for those seeking exposure to credit spreads.\n\n"
        "Evidence-Based Philosophy:\n"
        "Active management with a focus on deep credit analysis, differing from purely evidence-based approaches.\n\n"
        "Efficient Trading:\n"
        "Active trading in corporate bonds can incur higher costs but aims to capture credit opportunities.\n\n"
        "Track Record:\n"
        "Strong historical performance in capturing credit spreads, with a focus on corporate bonds.\n\n"
        "Recommendations:\n"
        "Excellent for clients seeking active management in corporate bonds with a focus on capturing credit opportunities."
    )
}

# ---------------------------
# Build the app interface
# ---------------------------
st.title("KYP Analysis Tool")

# Section 1: Fund Selection
st.header("1. Fund Selection")

st.subheader("Primary Fund Selection")
selected_equities = st.multiselect("Select Equities Funds:", options=list(equities_texts.keys()))
selected_fixed_income = st.multiselect("Select Fixed Income Funds:", options=list(fixed_income_texts.keys()))

st.subheader("Fund Comparison")
compare_equities = st.multiselect("Select Equities Funds for Comparison:", options=list(equities_texts.keys()))
compare_fixed_income = st.multiselect("Select Fixed Income Funds for Comparison:", options=list(fixed_income_texts.keys()))

# Section 2: Risk Evaluation Framework
st.header("2. Risk Evaluation Framework")

# Need to Take Risk
st.subheader("Need to Take Risk (Financial Need for Growth)")
col1, col2 = st.columns([1, 3])
with col1:
    risk_need = st.radio("Select your assessment for Need to Take Risk:", options=["High", "Moderate", "Low"])
with col2:
    st.markdown("""
    **Advisor Notes:**
    - Does the client need higher returns to meet their financial goals?
    - What is their required return to achieve financial goals (retirement, wealth accumulation)?
    - Do they have guaranteed income (pension, CPP, OAS, annuities)?
    - How flexible is their spending (can they reduce expenses if needed)?
    - Do they prioritize wealth accumulation or capital preservation?
    """)
risk_need_notes = st.text_area("Additional notes for Need to Take Risk:")

# Ability to Take Risk
st.subheader("Ability to Take Risk (Time Horizon & Financial Stability)")
col3, col4 = st.columns([1, 3])
with col3:
    risk_ability = st.radio("Select your assessment for Ability to Take Risk:", options=["High", "Moderate", "Low"])
with col4:
    st.markdown("""
    **Advisor Notes:**
    - What is the client's investment time horizon?
    - Will they rely on portfolio withdrawals soon?
    - Do they have liquidity needs?
    - How stable are other income sources?
    """)
risk_ability_notes = st.text_area("Additional notes for Ability to Take Risk:")

# Willingness to Take Risk
st.subheader("Willingness to Take Risk (Behavioral & Emotional Tolerance)")
col5, col6 = st.columns([1, 3])
with col5:
    risk_willingness = st.radio("Select your assessment for Willingness to Take Risk:", options=["High", "Moderate", "Low"])
with col6:
    st.markdown("""
    **Advisor Notes:**
    - How did the client react to past market downturns?
    - What is their investment experience and knowledge level?
    - How comfortable are they with volatility?
    - What are their expectations regarding risk vs. return?
    - Do they prioritize stability or maximizing returns?
    """)
risk_willingness_notes = st.text_area("Additional notes for Willingness to Take Risk:")

# Final Risk Profile Recommendation
st.subheader("Final Risk Profile Recommendation")
st.markdown("""
*(Based on the lowest score among Need, Ability, and Willingness.)*

**Assessment – Check One**

- **Aggressive (High Risk Tolerance):** High scores across all three categories.
- **Balanced (Moderate Risk Tolerance):** Moderate ability or willingness but high need.
- **Conservative (Low Risk Tolerance):** Low willingness or ability, regardless of need.
- **Ultra-Conservative (Minimal Risk):** Low ability and low willingness, even if higher returns are needed.
""")
final_risk_profile = st.radio("Select your Final Risk Profile Recommendation:", 
                              options=["Aggressive", "Balanced", "Conservative", "Ultra-Conservative"])
risk_conclusion = st.text_area("Optional final notes for Risk Evaluation:")

# Section 3: Client-Specific Recommendation
st.header("3. Client-Specific Recommendation")

client_name = st.text_input("Client Name", "Xavier")
investment_goals = st.text_area("Investment Goals", "Long-term growth, able to stomach market fluctuations.")
risk_tolerance = st.text_input("Risk Tolerance", "High")
account_type = st.text_input("Account Type", "TFSA")
primary_fund_recommendation = st.text_input("Primary Fund Recommended", "DFA Global Equity Portfolio F (DFA607)")
recommendation_notes = st.text_area("Recommendation Notes", "The primary recommendation is based on the client's preference for evidence-based, low-cost solutions.")

# ---------------------------
# Generate the KYP Analysis Report and Create a DOCX
# ---------------------------
if st.button("Generate KYP Analysis"):
    report_date = datetime.now().strftime("%Y-%m-%d")
    
    # Create a new Word document
    doc = docx.Document()
    
    # Set the default style font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title & Date
    doc.add_heading("KYP Analysis Report", level=1)
    doc.add_paragraph(f"Date: {report_date}")
    doc.add_paragraph("")
    
    # Section 1: Fund Selection
    doc.add_heading("1. Fund Selection", level=2)
    doc.add_heading("Primary Fund Selection", level=3)
    if selected_equities:
        doc.add_paragraph("Equities:", style='List Bullet')
        for fund in selected_equities:
            doc.add_paragraph(f"{fund}:", style='List Bullet 2')
            doc.add_paragraph(equities_texts[fund], style='Normal')
    if selected_fixed_income:
        doc.add_paragraph("Fixed Income:", style='List Bullet')
        for fund in selected_fixed_income:
            doc.add_paragraph(f"{fund}:", style='List Bullet 2')
            doc.add_paragraph(fixed_income_texts[fund], style='Normal')
    
    doc.add_heading("Fund Comparison", level=3)
    if compare_equities:
        doc.add_paragraph("Equities Comparison:", style='List Bullet')
        for fund in compare_equities:
            doc.add_paragraph(f"{fund}:", style='List Bullet 2')
            doc.add_paragraph(equities_texts[fund], style='Normal')
    if compare_fixed_income:
        doc.add_paragraph("Fixed Income Comparison:", style='List Bullet')
        for fund in compare_fixed_income:
            doc.add_paragraph(f"{fund}:", style='List Bullet 2')
            doc.add_paragraph(fixed_income_texts[fund], style='Normal')
    
    # Section 2: Risk Evaluation Framework
    doc.add_heading("2. Risk Evaluation Framework", level=2)
    
    doc.add_heading("Need to Take Risk (Financial Need for Growth)", level=3)
    doc.add_paragraph(f"Assessment: {risk_need}")
    doc.add_paragraph("Advisor Notes:")
    doc.add_paragraph(
        "- Does the client need higher returns to meet their financial goals?\n"
        "- What is their required return to achieve financial goals (retirement, wealth accumulation)?\n"
        "- Do they have guaranteed income (pension, CPP, OAS, annuities)?\n"
        "- How flexible is their spending (can they reduce expenses if needed)?\n"
        "- Do they prioritize wealth accumulation or capital preservation?",
        style='List Bullet'
    )
    doc.add_paragraph(f"Additional Notes: {risk_need_notes}")
    
    doc.add_heading("Ability to Take Risk (Time Horizon & Financial Stability)", level=3)
    doc.add_paragraph(f"Assessment: {risk_ability}")
    doc.add_paragraph("Advisor Notes:")
    doc.add_paragraph(
        "- What is the client's investment time horizon?\n"
        "- Will they rely on portfolio withdrawals soon?\n"
        "- Do they have liquidity needs?\n"
        "- How stable are other income sources?",
        style='List Bullet'
    )
    doc.add_paragraph(f"Additional Notes: {risk_ability_notes}")
    
    doc.add_heading("Willingness to Take Risk (Behavioral & Emotional Tolerance)", level=3)
    doc.add_paragraph(f"Assessment: {risk_willingness}")
    doc.add_paragraph("Advisor Notes:")
    doc.add_paragraph(
        "- How did the client react to past market downturns?\n"
        "- What is their investment experience and knowledge level?\n"
        "- How comfortable are they with volatility?\n"
        "- What are their expectations regarding risk vs. return?\n"
        "- Do they prioritize stability or maximizing returns?",
        style='List Bullet'
    )
    doc.add_paragraph(f"Additional Notes: {risk_willingness_notes}")
    
    doc.add_heading("Final Risk Profile Recommendation", level=3)
    doc.add_paragraph(f"Assessment: {final_risk_profile}")
    doc.add_paragraph(f"Final Notes: {risk_conclusion}")
    
    # Section 3: Client-Specific Recommendation
    doc.add_heading("3. Client-Specific Recommendation", level=2)
    doc.add_paragraph(f"Client Name: {client_name}")
    doc.add_paragraph(f"Investment Goals: {investment_goals}")
    doc.add_paragraph(f"Risk Tolerance: {risk_tolerance}")
    doc.add_paragraph(f"Account Type: {account_type}")
    doc.add_paragraph(f"Primary Fund Recommended: {primary_fund_recommendation}")
    doc.add_paragraph(f"Recommendation Notes: {recommendation_notes}")
    
    # Save the document to an in-memory buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    st.subheader("Generated KYP Analysis Report")
    st.text("The Word document has been generated. Use the download button below.")
    
    # Download button for the Word document
    st.download_button(
        label="Download as Word Document",
        data=buffer.getvalue(),
        file_name="KYP_Analysis_Report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
